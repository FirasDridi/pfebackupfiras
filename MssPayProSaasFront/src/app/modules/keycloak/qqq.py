import os
from docx import Document
from docx.shared import Pt

# Set the project directory and output file path
project_directory = r'C:\workspaces\projet\MssPayProSaasFront\src'
output_file = r'C:\Users\FIRAS\Desktop\angular_project_code.docx'
included_extensions = ['.ts', '.html', '.scss', '.css']
excluded_files = [
    'main.ts', 'polyfills.ts', 'test.ts', 'environment.ts', 'angular.json',
    'tsconfig.app.json', 'tsconfig.json', 'tsconfig.spec.json', 'karma.conf.js',
    'package.json', 'README.md', '.gitignore'
]
included_specific_files = ['app.module.ts', 'app-routing.module.ts']

# Function to read file content
def read_file_content(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as infile:
            return infile.read()
    except Exception:
        try:
            with open(file_path, 'r', encoding='latin-1') as infile:
                return infile.read()
        except Exception as e:
            print(f"Failed to read {file_path}: {e}")
            return ''

document = Document()
style = document.styles['Normal']
font = style.font
font.name, font.size = 'Courier New', Pt(8)

# Process files in the project directory
for root, _, files in os.walk(project_directory):
    for file in files:
        if (any(file.endswith(ext) for ext in included_extensions) and file not in excluded_files) or file in included_specific_files:
            file_path = os.path.join(root, file)
            print(f"Processing: {file_path}")
            document.add_heading(f"{file}", level=2)
            document.add_paragraph(f"Path: {file_path}")
            document.add_paragraph(read_file_content(file_path))

document.save(output_file)
print(f"Files saved to {output_file}")
